import os
import io
import json
from datetime import datetime
from typing import List, Dict, Tuple

import streamlit as st
import pandas as pd
import fitz  # PyMuPDF
from docx import Document as DocxDocument
from openai import OpenAI

# Optional OCR imports
try:
    import pytesseract
    from PIL import Image
    OCR_AVAILABLE = True
except Exception:
    OCR_AVAILABLE = False

# =============================
# Caching & Utilities
# =============================

@st.cache_data(show_spinner=False)
def _hash_bytes(b: bytes) -> str:
    return str(sum(b) % 10_000_000)

# =============================
# Extraction helpers (cached)
# =============================

@st.cache_data(show_spinner=False)
def extract_text_from_pdf_bytes(data: bytes, enable_ocr: bool = False, ocr_min_chars: int = 50) -> Tuple[str, List[int], Dict[int, int]]:
    """Return (full_text, page_starts, ocr_stats).
    If enable_ocr is True and a page appears mostly image (few chars), attempt OCR.
    ocr_stats maps page_index -> added_chars from OCR.
    """
    _ = _hash_bytes(data)
    text_parts = []
    page_starts = []
    ocr_stats: Dict[int, int] = {}
    with fitz.open(stream=data, filetype="pdf") as doc:
        for i, page in enumerate(doc):
            page_starts.append(sum(len(t) for t in text_parts))
            page_text = page.get_text()
            if enable_ocr and OCR_AVAILABLE and len(page_text.strip()) < ocr_min_chars:
                try:
                    pix = page.get_pixmap(dpi=200)
                    mode = "RGB" if not pix.alpha else "RGBA"
                    img = Image.frombytes(mode, [pix.width, pix.height], pix.samples)
                    ocr_text = pytesseract.image_to_string(img)
                    if ocr_text and len(ocr_text.strip()) > 0:
                        page_text = (page_text + "\n" + ocr_text).strip()
                        ocr_stats[i] = len(ocr_text)
                except Exception:
                    pass
            text_parts.append(page_text)
    full_text = "\n".join(text_parts)
    return full_text, page_starts, ocr_stats

@st.cache_data(show_spinner=False)
def extract_text_from_docx_bytes(data: bytes) -> str:
    _ = _hash_bytes(data)
    doc = DocxDocument(io.BytesIO(data))
    parts = []
    for para in doc.paragraphs:
        if para.text:
            parts.append(para.text)
    for table in doc.tables:
        for row in table.rows:
            parts.append("\t".join([cell.text for cell in row.cells]))
    return "\n".join(parts)

@st.cache_data(show_spinner=False)
def chunk_text(text: str, max_chars: int = 12000, overlap: int = 600) -> List[str]:
    if len(text) <= max_chars:
        return [text]
    chunks = []
    start = 0
    while start < len(text):
        end = min(start + max_chars, len(text))
        chunks.append(text[start:end])
        start = end - overlap
        if start < 0:
            start = 0
        if start >= len(text):
            break
    return chunks

# =============================
# Prompting
# =============================

BASE_SYSTEM_PROMPT = """
You are an expert contracts analyst specializing in Requests for Proposal (RFPs).
You read raw RFP text and produce four structured sections:
1) instructions: A clear, step-by-step submission checklist (with due dates, delivery portals, file formats, required forms, certifications, page limits, font/formatting requirements, questions deadlines, and evaluation criteria if present).
2) risks: Key contractual risk areas with severity (Low/Medium/High), short rationale, and mitigation. Include page references/snippets when available.
3) key_terms: Mentions of user-specified key terms/clauses (if provided) with a short excerpt and why they matter.
4) requirements: A compliance matrix of explicit requirements ("Requirement", "Citation", "Page", "Notes"). Keep Notes brief; set "Status" to "TBD" for the vendor to complete later.

Context profile: <<CONTEXT_PROFILE>>
- If Government (U.S. Federal), focus on FAR/DFARS references, proposal volumes, forms (SF 1449/33, SF1449, SF33), reps & certs (SAM, Section 889), submission portals (SAM, PIEE), compliance dates, and typical Gov risks (Termination for Convenience, data rights, MFC, audit).
- If Commercial/Non-government, focus on indemnities, limitation of liability caps, IP ownership, SLAs/LDs, payment terms, data security/privacy, insurance, jurisdiction/venue.

Output strictly as JSON with the schema:
{
  "instructions": [
    {"step": int, "item": str, "details": str, "due_date": str|null, "submission": str|null, "format": str|null, "page_limit": str|null, "reference": {"page": int|null, "snippet": str|null}}
  ],
  "risks": [
    {"topic": str, "severity": "Low"|"Medium"|"High", "why_risky": str, "mitigation": str, "reference": {"page": int|null, "snippet": str|null}}
  ],
  "key_terms": [
    {"term": str, "found": bool, "context": str|null, "reference": {"page": int|null, "snippet": str|null}}
  ],
  "requirements": [
    {"requirement": str, "citation": str|null, "page": int|null, "status": "TBD", "notes": str|null}
  ]
}
Use ISO dates when explicit; otherwise null. Keep entries concise but specific.
"""

def build_system_prompt(profile: str) -> str:
    profile_note = (
        "Government (U.S. Federal)" if profile.startswith("Government") else
        ("State/Local/Education" if profile.startswith("State/") else "Commercial / Non-government")
    )
    return BASE_SYSTEM_PROMPT.replace("<<CONTEXT_PROFILE>>", profile_note)

# =============================
# OpenAI calls
# =============================

def call_openai_analyze(client: OpenAI, model: str, system_prompt: str, rfp_text: str, key_terms: List[str]) -> Dict:
    user_payload = {
        "key_terms": key_terms,
        "rfp_excerpt": rfp_text[:1000] + ("..." if len(rfp_text) > 1000 else ""),
        "rfp_full": rfp_text
    }
    try:
        resp = client.responses.create(
            model=model,
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": json.dumps(user_payload)}
            ],
            response_format={"type": "json_object"}
        )
        if hasattr(resp, "output_text") and resp.output_text:
            content = resp.output_text
        elif hasattr(resp, "output") and resp.output:
            content = resp.output[0].content[0].text
        else:
            content = resp.choices[0].message.content
    except Exception:
        chat = client.chat.completions.create(
            model=model,
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": json.dumps(user_payload)}
            ],
            response_format={"type": "json_object"}
        )
        content = chat.choices[0].message.content
    try:
        data = json.loads(content)
    except Exception:
        data = {"instructions": [], "risks": [], "key_terms": [], "requirements": []}
    for k in ("instructions", "risks", "key_terms", "requirements"):
        data.setdefault(k, [])
    return data


def call_openai_questions(client: OpenAI, model: str, profile: str, extracted: Dict) -> List[str]:
    """Generate a concise list of clarifying questions for the issuer based on extracted findings."""
    sys = f"You are a senior proposal manager for {profile}. Generate concise, actionable clarification questions (max 12) based on gaps/ambiguities/risks in the extracted findings. Avoid yes/no; prefer specific asks with citations/pages when possible."
    user = json.dumps({
        "instructions": extracted.get("instructions", [])[:50],
        "risks": extracted.get("risks", [])[:50],
        "requirements": extracted.get("requirements", [])[:200]
    })
    try:
        resp = client.responses.create(
            model=model,
            messages=[{"role":"system","content":sys},{"role":"user","content":user}],
            response_format={"type":"json_object"}
        )
        text = resp.output_text if hasattr(resp, "output_text") else resp.output[0].content[0].text
        data = json.loads(text)
        # expect format {"questions": ["..."]}
        questions = data.get("questions") or data.get("items") or []
        if isinstance(questions, list):
            return [str(q) for q in questions][:12]
    except Exception:
        pass
    return []

# =============================
# UI
# =============================

st.set_page_config(page_title="RFP Analyzer", page_icon="📄", layout="wide")

st.title("📄 RFP Analyzer")
st.caption("Queue multiple RFPs (PDF/DOCX), then analyze them in one go. Get instructions, risks (with severity), key terms, a compliance matrix, and issuer questions. View HTML or download combined Excel.")

# --- Session state ---
if "batch_files" not in st.session_state:
    st.session_state.batch_files = []  # list of dicts {name, type, data}
if "uploader_key" not in st.session_state:
    st.session_state.uploader_key = "uploader-1"
if "last_settings" not in st.session_state:
    st.session_state.last_settings = {}

with st.sidebar:
    st.header("Settings")
    model = st.text_input("OpenAI model", value=os.getenv("OPENAI_MODEL", "gpt-4o-mini"))
    profile = st.radio(
        "RFP profile",
        ["Government (U.S. Federal)", "Commercial / Non-government", "State/Local/Education"],
        help="Tailors analysis for FAR/DFARS federal vs commercial vs state/local/ed nuances",
    )
    chunk_size = st.number_input("Chunk size (chars)", min_value=4000, max_value=20000, value=12000, step=500)
    overlap = st.number_input("Chunk overlap (chars)", min_value=0, max_value=3000, value=600, step=100)
    enable_ocr = st.toggle("Enable OCR for scanned PDFs (experimental)", value=False, help="Requires pytesseract & Pillow. Used only on pages with little text.")
    diagnostics = st.toggle("Diagnostics mode", value=False, help="Show tracebacks on errors")

    st.markdown("**Auth:** Use `st.secrets['OPENAI_API_KEY']` or env `OPENAI_API_KEY`.")

    # Optional: cost estimator
    with st.expander("Cost estimator (optional)"):
        input_cost = st.number_input("Input $/1K tokens", value=0.005, step=0.001, format="%.4f")
        output_cost = st.number_input("Output $/1K tokens", value=0.015, step=0.001, format="%.4f")

    # Save/Load settings
    with st.expander("Save/Load Settings"):
        if st.button("💾 Download current settings", use_container_width=True):
            cfg = {
                "model": model,
                "profile": profile,
                "chunk_size": int(chunk_size),
                "overlap": int(overlap),
                "enable_ocr": enable_ocr,
            }
            st.session_state.last_settings = cfg
            st.download_button("Download settings.json", data=json.dumps(cfg).encode("utf-8"), file_name="rfp_analyzer_settings.json", mime="application/json", use_container_width=True)
        cfg_file = st.file_uploader("Load settings.json", type=["json"], key="settings_uploader")
        if cfg_file is not None:
            try:
                cfg = json.loads(cfg_file.read())
                model = cfg.get("model", model)
                profile = cfg.get("profile", profile)
                chunk_size = cfg.get("chunk_size", chunk_size)
                overlap = cfg.get("overlap", overlap)
                enable_ocr = cfg.get("enable_ocr", enable_ocr)
                st.success("Settings loaded. They will apply to this session.")
            except Exception as e:
                st.error(f"Failed to load settings: {e}")

key_terms_input = st.text_area(
    "Optional: key terms/clauses to flag (comma-separated)",
    help="e.g., indemnification, limitation of liability, IP ownership, data rights, Section 889, audit rights"
)
key_terms = [t.strip() for t in key_terms_input.split(",") if t.strip()] if key_terms_input else []

# Upload CSV of extra key terms
with st.expander("Upload key terms CSV (column: term)"):
    csv_terms = st.file_uploader("CSV with 'term' column", type=["csv"], key="terms_csv")
    if csv_terms is not None:
        try:
            df_terms = pd.read_csv(csv_terms)
            extra = [str(x).strip() for x in df_terms.get("term", []) if str(x).strip()]
            key_terms = sorted(list({*key_terms, *extra}))
            st.success(f"Loaded {len(extra)} additional terms. Total terms: {len(key_terms)}")
        except Exception as e:
            st.error(f"Failed to parse CSV: {e}")

# --- File Staging ---
st.subheader("1) Add files to your batch")
uploaded_files = st.file_uploader(
    "Upload RFP files (PDF or DOCX)",
    type=["pdf", "docx"],
    accept_multiple_files=True,
    key=st.session_state.uploader_key,
)
col_add, col_clear = st.columns([1,1])
with col_add:
    if st.button("➕ Add selected to batch", use_container_width=True):
        if uploaded_files:
            for uf in uploaded_files:
                st.session_state.batch_files.append({
                    "name": uf.name,
                    "type": uf.name.lower().split(".")[-1],
                    "data": uf.getvalue(),
                    "approx_tokens": int(len(uf.getvalue()) / 4)
                })
            st.session_state.uploader_key = f"uploader-{len(st.session_state.batch_files)}"
            st.rerun()
with col_clear:
    if st.button("🗑️ Clear batch", use_container_width=True):
        st.session_state.batch_files = []
        st.session_state.uploader_key = "uploader-1"
        st.rerun()

if st.session_state.batch_files:
    st.success(f"Queued {len(st.session_state.batch_files)} file(s)")
    df_preview = pd.DataFrame([
        {"File": f["name"], "Type": f["type"], "Size (KB)": round(len(f["data"]) / 1024, 1)} for f in st.session_state.batch_files
    ])
    st.dataframe(df_preview, use_container_width=True)
else:
    st.info("No files in batch yet. Upload and click 'Add selected to batch'.")

# --- Analyze Button ---
st.subheader("2) Analyze")
run_analysis = st.button("▶️ Analyze batch", type="primary", use_container_width=True, disabled=not st.session_state.batch_files)

if run_analysis:
    # Setup OpenAI client
    api_key = None
    try:
        api_key = st.secrets.get("OPENAI_API_KEY")  # type: ignore
    except Exception:
        api_key = None
    if not api_key:
        api_key = os.getenv("OPENAI_API_KEY")
    if not api_key:
        st.error("No OpenAI API key found.")
        st.stop()

    client = OpenAI(api_key=api_key)
    system_prompt = build_system_prompt(profile)

    # Collect combined results across files
    combined_instr = []
    combined_risks = []
    combined_terms = []
    combined_reqs = []

    combined_html_parts = ["<h1>RFP Combined Analysis</h1>"]

    for f in st.session_state.batch_files:
        try:
            st.subheader(f"File: {f['name']}")
            ext = f["type"]
            if ext == "pdf":
                raw_text, _, ocr_stats = extract_text_from_pdf_bytes(f["data"], enable_ocr=enable_ocr)
                if enable_ocr and OCR_AVAILABLE and ocr_stats:
                    st.caption(f"OCR applied on {len(ocr_stats)} page(s): +{sum(ocr_stats.values())} chars")
                elif enable_ocr and not OCR_AVAILABLE:
                    st.warning("OCR not available (install Pillow & pytesseract). Proceeding without OCR.")
            elif ext == "docx":
                raw_text = extract_text_from_docx_bytes(f["data"]) 
            else:
                st.error("Unsupported file type.")
                continue

            chunks = chunk_text(raw_text, max_chars=int(chunk_size), overlap=int(overlap))
            st.caption(f"Parsed ~{len(raw_text):,} characters → {len(chunks)} chunk(s)")

            merged = {"instructions": [], "risks": [], "key_terms": [], "requirements": []}
            for idx, ch in enumerate(chunks, start=1):
                try:
                    with st.status(f"Analyzing chunk {idx}/{len(chunks)}…"):
                        result = call_openai_analyze(client, model, system_prompt, ch, key_terms)
                        for k in merged:
                            merged[k].extend(result.get(k, []))
                except Exception as e:
                    if diagnostics:
                        import traceback
                        st.exception(e)
                        st.text(traceback.format_exc())
                    else:
                        st.error(f"LLM analysis failed on chunk {idx}.")

            # DataFrames (per file)
            instr_df = pd.DataFrame(merged["instructions"]) if merged["instructions"] else pd.DataFrame(columns=["step","item","details","due_date","submission","format","page_limit","reference"]) 
            risks_df = pd.DataFrame(merged["risks"]) if merged["risks"] else pd.DataFrame(columns=["topic","severity","why_risky","mitigation","reference"]) 
            terms_df = pd.DataFrame(merged["key_terms"]) if merged["key_terms"] else pd.DataFrame(columns=["term","found","context","reference"]) 
            reqs_df = pd.DataFrame(merged["requirements"]) if merged["requirements"] else pd.DataFrame(columns=["requirement","citation","page","status","notes"]) 

            # Add File column for combined export
            for df in (instr_df, risks_df, terms_df, reqs_df):
                if not df.empty:
                    df.insert(0, "File", f['name'])

            # Append to combined lists
            if not instr_df.empty:
                combined_instr.append(instr_df)
            if not risks_df.empty:
                combined_risks.append(risks_df)
            if not terms_df.empty:
                combined_terms.append(terms_df)
            if not reqs_df.empty:
                combined_reqs.append(reqs_df)

            # Per-file HTML view + download
            with st.expander("📄 View results (HTML)", expanded=True):
                parts = [
                    f"<h2>{f['name']}</h2>",
                    "<h3>Submission Instructions</h3>",
                    instr_df.to_html(index=False, escape=False) if not instr_df.empty else "<p>No instructions found.</p>",
                    "<h3>Risk Areas</h3>",
                    risks_df.to_html(index=False, escape=False) if not risks_df.empty else "<p>No risks found.</p>",
                    "<h3>Key Terms</h3>",
                    terms_df.to_html(index=False, escape=False) if not terms_df.empty else "<p>No key term matches found.</p>",
                    "<h3>Compliance Matrix</h3>",
                    reqs_df.to_html(index=False, escape=False) if not reqs_df.empty else "<p>No explicit requirements extracted.</p>",
                ]
                html_bytes = "\n".join(parts).encode("utf-8")
                st.download_button(
                    label="⬇️ Download HTML (this file)",
                    data=html_bytes,
                    file_name=f"{os.path.splitext(f['name'])[0]}_analysis.html",
                    mime="text/html",
                    use_container_width=True,
                )
                combined_html_parts.extend(parts)

            # Export Excel per file
            output = io.BytesIO()
            with pd.ExcelWriter(output, engine="xlsxwriter") as writer:
                instr_df.to_excel(writer, sheet_name="Instructions", index=False)
                risks_df.to_excel(writer, sheet_name="Risks", index=False)
                terms_df.to_excel(writer, sheet_name="KeyTerms", index=False)
                reqs_df.to_excel(writer, sheet_name="Compliance", index=False)
                meta_df = pd.DataFrame({
                    "File": [f['name']],
                    "AnalyzedAt": [datetime.utcnow().strftime("%Y-%m-%d %H:%M:%SZ")],
                    "Model": [model],
                    "Profile": [profile],
                    "Chars": [len(raw_text)],
                    "Chunks": [len(chunks)]
                })
                meta_df.to_excel(writer, sheet_name="Meta", index=False)
            output.seek(0)

            st.download_button(
                "⬇️ Download Excel (this file)",
                output,
                f"{os.path.splitext(f['name'])[0]}_analysis.xlsx",
                "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                use_container_width=True,
            )
        except Exception as e:
            if diagnostics:
                import traceback
                st.exception(e)
                st.text(traceback.format_exc())
            else:
                st.error(f"Something went wrong processing {f['name']}.")

    # --- Combined exports across all files ---
    if combined_instr or combined_risks or combined_terms or combined_reqs:
        st.markdown("---")
        st.subheader("📦 Combined results (all files)")

        comb_instr_df = pd.concat(combined_instr, ignore_index=True) if combined_instr else pd.DataFrame(columns=["File","step","item","details","due_date","submission","format","page_limit","reference"]) 
        comb_risks_df = pd.concat(combined_risks, ignore_index=True) if combined_risks else pd.DataFrame(columns=["File","topic","severity","why_risky","mitigation","reference"]) 
        comb_terms_df = pd.concat(combined_terms, ignore_index=True) if combined_terms else pd.DataFrame(columns=["File","term","found","context","reference"]) 
        comb_reqs_df = pd.concat(combined_reqs, ignore_index=True) if combined_reqs else pd.DataFrame(columns=["File","requirement","citation","page","status","notes"]) 

        # ---- Inline filters ----
        with st.expander("Filters", expanded=False):
            files = sorted({*comb_instr_df.get("File", pd.Series(dtype=str)).unique(), *comb_risks_df.get("File", pd.Series(dtype=str)).unique(), *comb_terms_df.get("File", pd.Series(dtype=str)).unique(), *comb_reqs_df.get("File", pd.Series(dtype=str)).unique()})
            sel_files = st.multiselect("Filter by file", files, default=files)
            sev_vals = ["High", "Medium", "Low"]
            sel_sev = st.multiselect("Risk severity", sev_vals, default=sev_vals)
            text_q = st.text_input("Search text (applies to visible previews)")

        def _apply_filters(df: pd.DataFrame) -> pd.DataFrame:
            if df.empty:
                return df
            if "File" in df.columns and sel_files:
                df = df[df["File"].isin(sel_files)]
            if set(["severity"]).issubset(df.columns) and sel_sev:
                df = df[df["severity"].isin(sel_sev)]
            if text_q:
                mask = pd.Series(False, index=df.index)
                for c in df.columns:
                    mask = mask | df[c].astype(str).str.contains(text_q, case=False, na=False)
                df = df[mask]
            return df

        p_instr = _apply_filters(comb_instr_df)
        p_risks = _apply_filters(comb_risks_df)
        p_terms = _apply_filters(comb_terms_df)
        p_reqs  = _apply_filters(comb_reqs_df)

        with st.expander("Preview combined tables", expanded=False):
            st.write("Instructions (first 200 rows)")
            st.dataframe(p_instr.head(200), use_container_width=True)
            st.write("Risks (first 200 rows)")
            st.dataframe(p_risks.head(200), use_container_width=True)
            st.write("Key Terms (first 200 rows)")
            st.dataframe(p_terms.head(200), use_container_width=True)
            st.write("Compliance (first 200 rows)")
            st.dataframe(p_reqs.head(200), use_container_width=True)

        # Combined Excel
        comb_xlsx = io.BytesIO()
        with pd.ExcelWriter(comb_xlsx, engine="xlsxwriter") as writer:
            comb_instr_df.to_excel(writer, sheet_name="Instructions_All", index=False)
            comb_risks_df.to_excel(writer, sheet_name="Risks_All", index=False)
            comb_terms_df.to_excel(writer, sheet_name="KeyTerms_All", index=False)
            comb_reqs_df.to_excel(writer, sheet_name="Compliance_All", index=False)
            meta = pd.DataFrame({
                "GeneratedAt": [datetime.utcnow().strftime("%Y-%m-%d %H:%M:%SZ")],
                "Model": [model],
                "Profile": [profile],
                "Files": [len(st.session_state.batch_files)]
            })
            meta.to_excel(writer, sheet_name="Meta", index=False)
        comb_xlsx.seek(0)
        st.download_button(
            "⬇️ Download Combined Excel (all files)",
            comb_xlsx,
            "rfp_analysis_combined.xlsx",
            "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            use_container_width=True,
        )

        # Combined HTML (stitched)
        comb_html_bytes = "\n".join(combined_html_parts).encode("utf-8")
        with st.expander("🌐 View combined HTML", expanded=False):
            st.components.v1.html("\n".join(combined_html_parts), height=600, scrolling=True)
        st.download_button(
            "⬇️ Download Combined HTML",
            comb_html_bytes,
            "rfp_analysis_combined.html",
            "text/html",
            use_container_width=True,
        )

        # ---- Calendar export (ICS) from instruction due dates ----
        if not comb_instr_df.empty and "due_date" in comb_instr_df.columns:
            try:
                ics = [
                    "BEGIN:VCALENDAR",
                    "VERSION:2.0",
                    "PRODID:-//RFP Analyzer//EN"
                ]
                for _, row in comb_instr_df.iterrows():
                    due = row.get("due_date")
                    if pd.isna(due) or not str(due).strip():
                        continue
                    try:
                        dt = pd.to_datetime(due).date()
                        uid = f"rfp-{hash(str(row))}@rfp-analyzer"
                        summary = f"RFP: {row.get('item','Submission step')}"
                        ics += [
                            "BEGIN:VEVENT",
                            f"UID:{uid}",
                            f"DTSTAMP:{datetime.utcnow().strftime('%Y%m%dT%H%M%SZ')}",
                            f"DTSTART;VALUE=DATE:{dt.strftime('%Y%m%d')}",
                            f"SUMMARY:{summary}",
                            f"DESCRIPTION:{(row.get('details') or '')[:300]}",
                            "END:VEVENT",
                        ]
                    except Exception:
                        continue
                ics.append("END:VCALENDAR")
                ics_bytes = "\n".join(ics).encode("utf-8")
                st.download_button("⬇️ Download Due Dates (.ics)", ics_bytes, "rfp_due_dates.ics", "text/calendar", use_container_width=True)
            except Exception as e:
                if diagnostics:
                    st.error(f"ICS generation failed: {e}")

        # ---- Vendor response starter kit (Excel) ----
        if not comb_reqs_df.empty:
            kit = io.BytesIO()
            kit_df = comb_reqs_df.copy()
            # Add management columns
            for col in ["Owner", "Due", "Response", "Evidence Link", "Status", "Notes"]:
                if col not in kit_df.columns:
                    kit_df[col] = "" if col != "Status" else "TBD"
            with pd.ExcelWriter(kit, engine="xlsxwriter") as writer:
                kit_df.to_excel(writer, sheet_name="Compliance Tracker", index=False)
                # simple README sheet
                readme = pd.DataFrame({
                    "RFP Analyzer Compliance Tracker": [
                        "Fill Owner/Due/Response per row.",
                        "Use Status values: TBD / In Progress / Complete / Risk.",
                        "Attach evidence links where applicable.",
                    ]
                })
                readme.to_excel(writer, sheet_name="README", index=False)
            kit.seek(0)
            st.download_button("⬇️ Download Response Starter Kit (Excel)", kit, "rfp_response_kit.xlsx", "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet", use_container_width=True)

        # ---- Clarifying questions for issuer ----
        extracted_sample = {
            "instructions": comb_instr_df.to_dict(orient="records")[:50],
            "risks": comb_risks_df.to_dict(orient="records")[:50],
            "requirements": comb_reqs_df.to_dict(orient="records")[:200],
        }
        qs = call_openai_questions(client, model, profile, extracted_sample)
        with st.expander("❓ Suggested questions for the issuer", expanded=False):
            if qs:
                for i, q in enumerate(qs, 1):
                    st.markdown(f"**Q{i}.** {q}")
            else:
                st.caption("No suggested questions generated (or model declined).")

    # Optional: token/cost estimate summary
    with st.expander("Est. token & cost summary", expanded=False):
        if st.session_state.batch_files:
            approx_in_tokens = sum(len(f["data"]) for f in st.session_state.batch_files) / 4
            approx_out_tokens = approx_in_tokens * 0.1
            st.write({
                "approx_input_tokens": int(approx_in_tokens),
                "approx_output_tokens": int(approx_out_tokens),
                "est_input_cost_usd": round((approx_in_tokens/1000)*input_cost, 4),
                "est_output_cost_usd": round((approx_out_tokens/1000)*output_cost, 4),
                "est_total_cost_usd": round((approx_in_tokens/1000)*input_cost + (approx_out_tokens/1000)*output_cost, 4)
            })

# Footer
st.markdown("---")
st.caption("Tip: Queue files in several rounds, then click Analyze once. Use filters & exports to build your compliance matrix, calendar, and risk register.")
